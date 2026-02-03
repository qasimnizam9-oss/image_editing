import os
import pandas as pd
from flask import Flask, render_template, request, flash, redirect, url_for, session
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from PIL import Image
import img2pdf
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
app.secret_key = "editing_monkey_secret_key"

# --- MYSQL CONFIGURATION ---
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root:qasimnizam123.@localhost/editing_monkey'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# Configuration for files
UPLOAD_FOLDER = 'uploads'
STATIC_FOLDER = 'static/processed'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif','jfif'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(STATIC_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# --- DATABASE MODELS ---
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    first_name = db.Column(db.String(50), nullable=False)
    last_name = db.Column(db.String(50), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)

class ContactMessage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), nullable=False)
    message = db.Column(db.Text, nullable=False)

# Create tables in MySQL
with app.app_context():
    db.create_all()

# --- HELPER FUNCTIONS ---
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def process_image(filename, operation):
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file_basename = filename.rsplit('.', 1)[0]
    
    if operation == "pdf":
        new_filename = f"{file_basename}.pdf"
        with open(os.path.join(STATIC_FOLDER, new_filename), "wb") as f:
            f.write(img2pdf.convert(input_path))
            
    elif operation == "docx":
        new_filename = f"{file_basename}.docx"
        doc = Document()
        doc.add_heading('Converted Image', 0)
        doc.add_picture(input_path, width=Inches(5)) 
        doc.save(os.path.join(STATIC_FOLDER, new_filename))

    elif operation in ["xlsx", "csv"]:
        new_filename = f"{file_basename}.{operation}"
        img = Image.open(input_path)
        df = pd.DataFrame({
            "Attribute": ["Filename", "Width", "Height"],
            "Value": [filename, img.width, img.height]
        })
        output_path = os.path.join(STATIC_FOLDER, new_filename)
        if operation == "xlsx": df.to_excel(output_path, index=False)
        else: df.to_csv(output_path, index=False)
        
    else:
        img = Image.open(input_path)
        if operation == "grayscale": img = img.convert('L')
        new_filename = f"{file_basename}_edited.png"
        img.save(os.path.join(STATIC_FOLDER, new_filename))
        
    return new_filename

# --- NAVIGATION ROUTES ---

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/how")
def how():
    return render_template("how.html")

@app.route("/contact", methods=["GET", "POST"])
def contact():
    if request.method == "POST":
        name = request.form.get("name")
        email = request.form.get("email")
        msg_content = request.form.get("message")
        
        # Save message to MySQL
        new_msg = ContactMessage(name=name, email=email, message=msg_content)
        db.session.add(new_msg)
        db.session.commit()
        
        flash(f"Thanks {name}! Your message has been saved in our database.")
        return redirect(url_for('contact'))
    return render_template("contact.html")

# --- AUTH ROUTES ---

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        hashed_pw = generate_password_hash(request.form.get("password"), method='pbkdf2:sha256')
        new_user = User(
            first_name=request.form.get("fname"),
            last_name=request.form.get("lname"),
            email=request.form.get("email"),
            password=hashed_pw
        )
        try:
            db.session.add(new_user)
            db.session.commit()
            flash("Account created! Please login.")
            return redirect(url_for('login'))
        except:
            flash("Email already exists!")
    return render_template("signup.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")
        user = User.query.filter_by(email=email).first()
        
        if user and check_password_hash(user.password, password):
            session['user_id'] = user.id
            session['user_name'] = user.first_name
            flash(f"Welcome back, {user.first_name}!")
            return redirect(url_for('home'))
        flash("Invalid login credentials.")
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for('home'))

# --- PROCESSING ROUTE ---

@app.route("/edit", methods=["POST"])
def edit():
    operation = request.form.get("operation")
    file = request.files.get('file')

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        processed_file = process_image(filename, operation)
        os.remove(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        return render_template("index.html", processed_file=processed_file)
    
    flash("Please upload a valid image.")
    return redirect(url_for('home'))

if __name__ == "__main__":
    app.run(debug=True, port=5001)